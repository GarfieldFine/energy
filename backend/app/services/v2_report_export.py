from __future__ import annotations

import io
from datetime import datetime
from typing import Any

import pandas as pd

from app.services.energy_store import load_energy
from app.services import stats_service
from app.services.v2_service import forecast_energy, ops_indicators, ops_suggestions

try:
    from docx import Document

    _HAS_DOCX = True
except ImportError:
    Document = None  # type: ignore[assignment]
    _HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    _HAS_REPORTLAB = True
except ImportError:
    A4 = None  # type: ignore[assignment]
    colors = None  # type: ignore[assignment]
    getSampleStyleSheet = None  # type: ignore[assignment]
    pdfmetrics = None  # type: ignore[assignment]
    TTFont = None  # type: ignore[assignment]
    Paragraph = None  # type: ignore[assignment]
    SimpleDocTemplate = None  # type: ignore[assignment]
    Spacer = None  # type: ignore[assignment]
    Table = None  # type: ignore[assignment]
    TableStyle = None  # type: ignore[assignment]
    _HAS_REPORTLAB = False


def _ensure_docx() -> bool:
    global _HAS_DOCX, Document
    if _HAS_DOCX:
        return True
    try:
        from docx import Document as _Document

        Document = _Document  # type: ignore[assignment]
        _HAS_DOCX = True
    except ImportError:
        _HAS_DOCX = False
    return _HAS_DOCX


def _ensure_reportlab() -> bool:
    global _HAS_REPORTLAB, A4, colors, getSampleStyleSheet, pdfmetrics, TTFont
    global Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    if _HAS_REPORTLAB:
        return True
    try:
        from reportlab.lib.pagesizes import A4 as _A4
        from reportlab.lib import colors as _colors
        from reportlab.lib.styles import getSampleStyleSheet as _getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics as _pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont as _TTFont
        from reportlab.platypus import (
            Paragraph as _Paragraph,
            SimpleDocTemplate as _SimpleDocTemplate,
            Spacer as _Spacer,
            Table as _Table,
            TableStyle as _TableStyle,
        )

        A4 = _A4  # type: ignore[assignment]
        colors = _colors  # type: ignore[assignment]
        getSampleStyleSheet = _getSampleStyleSheet  # type: ignore[assignment]
        pdfmetrics = _pdfmetrics  # type: ignore[assignment]
        TTFont = _TTFont  # type: ignore[assignment]
        Paragraph = _Paragraph  # type: ignore[assignment]
        SimpleDocTemplate = _SimpleDocTemplate  # type: ignore[assignment]
        Spacer = _Spacer  # type: ignore[assignment]
        Table = _Table  # type: ignore[assignment]
        TableStyle = _TableStyle  # type: ignore[assignment]
        _HAS_REPORTLAB = True
    except ImportError:
        _HAS_REPORTLAB = False
    return _HAS_REPORTLAB


def _fmt(v: Any, digits: int = 2) -> str:
    if v is None:
        return "-"
    try:
        x = float(v)
    except (TypeError, ValueError):
        return str(v)
    if pd.isna(x):
        return "-"
    return f"{x:.{digits}f}"


def _pick_buildings(building_id: str | None) -> list[str]:
    if building_id:
        return [building_id]
    df = load_energy().copy()
    if df.empty or "building_id" not in df.columns:
        return []
    if "electricity_kwh" not in df.columns:
        return [str(x) for x in df["building_id"].dropna().astype(str).unique()[:5]]
    df["electricity_kwh"] = pd.to_numeric(df["electricity_kwh"], errors="coerce")
    g = (
        df.dropna(subset=["building_id"])
        .groupby("building_id", as_index=False)["electricity_kwh"]
        .sum()
        .sort_values("electricity_kwh", ascending=False)
        .head(5)
    )
    return [str(x) for x in g["building_id"].tolist()]


def _building_pack(bid: str) -> dict[str, Any]:
    period = stats_service.period_summary(bid, None, None)
    anomaly = stats_service.anomaly_analysis(bid, None, None, z_threshold=3.0)
    ind = ops_indicators(bid).get("indicators", {})
    sug = ops_suggestions(bid).get("items", [])
    fc = forecast_energy(bid, horizon_hours=24)
    values = fc.get("values") or []
    next24_avg = float(sum(values) / len(values)) if values else None
    ewi = ind.get("ewi")
    dh = ind.get("dh")
    ratio = anomaly.get("ratio") or 0.0
    problems: list[str] = []
    if ratio >= 0.1:
        problems.append(f"异常小时占比偏高（{ratio*100:.2f}%）")
    elif ratio >= 0.05:
        problems.append(f"异常小时占比需关注（{ratio*100:.2f}%）")
    if ewi is not None and float(ewi) >= 0.7:
        problems.append("夜间基荷偏高，存在待机/空转风险")
    if dh is not None and float(dh) <= 0.75:
        problems.append("设备健康度偏低，需加强巡检与预防性维护")
    if not problems:
        problems.append("运行总体稳定，未见显著异常风险")
    return {
        "building_id": bid,
        "period": period,
        "anomaly": anomaly,
        "indicators": ind,
        "suggestions": sug,
        "forecast": {"model": fc.get("model"), "next24_avg_kwh_h": next24_avg},
        "problems": problems,
    }


def _report_context(kind: str, building_id: str | None) -> dict[str, Any]:
    bids = _pick_buildings(building_id)
    packs = [_building_pack(b) for b in bids]
    total_rows = sum(int((p.get("period") or {}).get("rows") or 0) for p in packs)
    return {
        "kind": kind,
        "building_id": building_id,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "scope_note": f"共 {len(bids)} 栋建筑（按总电耗排序）" if not building_id else "单栋建筑模式",
        "buildings": packs,
        "total_rows": total_rows,
    }


def _esg_summary_line(pack: dict[str, Any]) -> tuple[str, str, str]:
    p = pack.get("period") or {}
    a = pack.get("anomaly") or {}
    i = pack.get("indicators") or {}
    sums = p.get("sums") or {}
    e = f"电耗 { _fmt(sums.get('electricity_kwh')) } kWh，异常占比 { _fmt((a.get('ratio') or 0) * 100) }%"
    s = "舒适度建议：结合温湿度均值做空调时段策略优化（本版以运营数据近似）"
    g = f"EWI { _fmt(i.get('ewi'), 4) }，DH { _fmt(i.get('dh'), 4) }，建议闭环 {len(pack.get('suggestions') or [])} 条"
    return e, s, g


def _owner_due(priority: str | None) -> tuple[str, str]:
    p = (priority or "").lower()
    if p == "high":
        return "能源经理/设备主管", "3 个工作日内"
    if p == "medium":
        return "运维工程师", "7 个工作日内"
    return "值班运维", "15 个工作日内"


def _risk_level(pack: dict[str, Any]) -> str:
    a = pack.get("anomaly") or {}
    i = pack.get("indicators") or {}
    ratio = float(a.get("ratio") or 0.0)
    ewi = float(i.get("ewi") or 0.0)
    dh = float(i.get("dh") or 1.0)
    score = 0
    if ratio >= 0.1:
        score += 3
    elif ratio >= 0.05:
        score += 2
    elif ratio >= 0.02:
        score += 1
    if ewi >= 0.9:
        score += 3
    elif ewi >= 0.7:
        score += 2
    elif ewi >= 0.5:
        score += 1
    if dh <= 0.65:
        score += 3
    elif dh <= 0.75:
        score += 2
    elif dh <= 0.85:
        score += 1
    if score >= 7:
        return "高"
    if score >= 4:
        return "中"
    return "低"


def _action_plan_rows(pack: dict[str, Any]) -> list[list[str]]:
    out: list[list[str]] = []
    for r in _five_segment_rows(pack):
        out.append([r["priority"], r["action"], r["owner"], r["due"]])
    return out


def _five_segment_rows(pack: dict[str, Any]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    problems = pack.get("problems") or []
    suggestions = pack.get("suggestions") or []
    anomaly = pack.get("anomaly") or {}
    ratio = float(anomaly.get("ratio") or 0.0)
    for i, it in enumerate(suggestions, 1):
        pri = str(it.get("priority") or "-")
        owner, due = _owner_due(pri)
        phenomenon = problems[min(i - 1, len(problems) - 1)] if problems else "运行总体稳定，未见显著异常风险"
        impact = (
            f"若不处置，异常占比可能维持在 {ratio*100:.2f}% 左右，增加能耗与故障风险。"
            if ratio >= 0.05
            else "若不处置，将削弱节能稳定性并影响后续指标优化。"
        )
        title_txt = str(it.get("title") or "").strip()
        exp = (
            f"预计节电约 {_fmt(it.get('expected_saving_kwh_per_hour'))} kWh/h。"
            if it.get("expected_saving_kwh_per_hour") is not None
            else str(it.get("expected_effect") or "按周复核策略参数。")
        )
        rows.append(
            {
                "priority": pri,
                "phenomenon": phenomenon,
                "impact": impact,
                "action": f"{title_txt}（{exp}）",
                "owner": owner,
                "due": due,
            }
        )
    if not rows:
        rows.append(
            {
                "priority": "low",
                "phenomenon": "当前运行总体稳定。",
                "impact": "维持现状可满足当前演示与运行要求。",
                "action": "按月复核阈值、工单闭环时效与设备点检记录。",
                "owner": "运维工程师",
                "due": "每月",
            }
        )
    return rows


def _docx_add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v


def _build_word(ctx: dict[str, Any], kind: str) -> tuple[bytes, str, str]:
    bio = io.BytesIO()
    doc = Document()
    title = "建筑能源运营优化报告" if kind == "operations" else "建筑能源 ESG 专项报告"
    section = doc.sections[0]
    section.header.paragraphs[0].text = f"{title}（建筑能源智能管理系统）"
    section.footer.paragraphs[0].text = "建筑能源智能管理系统 | 自动生成报告 | 仅供内部评审"
    doc.add_heading(title, 0)
    doc.add_paragraph("项目名称：建筑能源智能管理系统")
    doc.add_paragraph("文档级别：内部评审版")
    doc.add_paragraph(f"生成时间：{ctx['generated_at']}")
    doc.add_paragraph(f"覆盖范围：{ctx['scope_note']}")
    if ctx.get("building_id"):
        doc.add_paragraph(f"目标建筑：{ctx['building_id']}")

    doc.add_heading("一、报告说明", level=1)
    doc.add_paragraph("数据来源：系统小时级能耗数据、异常检测结果与运营规则库。")
    doc.add_paragraph("分析范围：默认按总电耗选取前 5 栋建筑；若指定建筑 ID 则按单栋输出。")
    doc.add_paragraph("输出形式：结构化文字结论与整改建议，适用于评审与运维复盘。")
    doc.add_paragraph("口径说明：异常检测采用 z-score（阈值=3.0）；EWI 反映夜间基荷浪费倾向；DH 反映设备运行健康近似值。")
    doc.add_paragraph("说明：空间利用相关建议不在本报告输出，室内布局建议请以“孪生与视觉”页面上传建模结果为准。")

    doc.add_heading("二、建筑清单与核心指标总览", level=1)
    rows = []
    for p in ctx["buildings"]:
        period = p["period"]
        anomaly = p["anomaly"]
        ind = p["indicators"]
        rows.append(
            [
                str(p["building_id"]),
                _risk_level(p),
                str(period.get("rows", 0)),
                _fmt((period.get("sums") or {}).get("electricity_kwh")),
                _fmt((anomaly.get("ratio") or 0) * 100),
                _fmt(ind.get("ewi"), 4),
                _fmt(ind.get("dh"), 4),
                _fmt((p.get("forecast") or {}).get("next24_avg_kwh_h")),
            ]
        )
    _docx_add_table(
        doc,
        ["建筑", "风险等级", "样本行数", "累计电耗(kWh)", "异常占比(%)", "EWI", "DH", "未来24h均值(kWh/h)"],
        rows,
    )
    doc.add_paragraph("风险等级判定：综合异常占比、EWI、DH 三项指标，分为高/中/低。")

    sec_title = "三、逐栋异常问题报告与运营建议" if kind == "operations" else "三、逐栋 ESG 评估与整改建议"
    doc.add_heading(sec_title, level=1)
    for idx, p in enumerate(ctx["buildings"], 1):
        bid = p["building_id"]
        period = p.get("period") or {}
        sums = period.get("sums") or {}
        means = period.get("means") or {}
        tr = period.get("time_range") or {}
        anomaly = p.get("anomaly") or {}
        indicators = p.get("indicators") or {}
        fc = p.get("forecast") or {}
        doc.add_heading(f"{idx}. 建筑 {bid}", level=2)
        doc.add_paragraph("监测概况：")
        doc.add_paragraph(
            f"• 样本行数：{period.get('rows', 0)}；时间范围：{tr.get('min', '-')[:19]} 至 {tr.get('max', '-')[:19]}。"
        )
        doc.add_paragraph(
            f"• 累计电耗：{_fmt(sums.get('electricity_kwh'))} kWh；平均小时电耗：{_fmt(means.get('electricity_kwh'))} kWh/h。"
        )
        doc.add_paragraph(
            f"• 异常统计：异常小时 {anomaly.get('anomaly_hours', 0)} / 总小时 {anomaly.get('total_hours', 0)}，"
            f"占比 {_fmt((anomaly.get('ratio') or 0) * 100)}%。"
        )
        doc.add_paragraph(
            f"• 指标：EWI={_fmt(indicators.get('ewi'),4)}，DH={_fmt(indicators.get('dh'),4)}；"
            f"未来24h均值预测={_fmt(fc.get('next24_avg_kwh_h'))} kWh/h（模型：{fc.get('model') or '-'}）。"
        )
        doc.add_paragraph(f"• 综合风险等级：{_risk_level(p)}。")
        doc.add_paragraph("异常问题：")
        for x in p["problems"]:
            doc.add_paragraph(f"• {x}")
        samples = (anomaly.get("samples") or [])[:5]
        doc.add_paragraph("异常样本（前 5 条）：")
        if samples:
            s_rows = []
            for s in samples:
                s_rows.append(
                    [
                        str(s.get("monitor_time", "-")),
                        _fmt(s.get("electricity_kwh")),
                        str(s.get("building_id", bid)),
                    ]
                )
            _docx_add_table(doc, ["监测时间", "市电(kWh)", "建筑"], s_rows)
        else:
            doc.add_paragraph("• 当前阈值下无异常样本。")
        doc.add_paragraph("五段式整改建议（现象-影响-处置-责任人-时限）：")
        for j, r in enumerate(_five_segment_rows(p), 1):
            doc.add_paragraph(f"{j}) [{r['priority']}] 现象：{r['phenomenon']}")
            doc.add_paragraph(f"    影响：{r['impact']}")
            doc.add_paragraph(f"    处置：{r['action']}")
            doc.add_paragraph(f"    责任人：{r['owner']}；完成时限：{r['due']}")
        doc.add_paragraph("整改计划表：")
        _docx_add_table(doc, ["优先级", "措施", "责任人", "完成时限"], _action_plan_rows(p))
        if kind == "esg":
            e, s, g = _esg_summary_line(p)
            doc.add_paragraph(f"E（环境）：{e}")
            doc.add_paragraph(f"S（社会）：{s}")
            doc.add_paragraph(f"G（治理）：{g}")
            doc.add_paragraph("ESG专项建议：优先落实夜间基荷治理与设备巡检闭环，持续跟踪月度指标波动并形成台账。")

    tail_title = "四、结论与实施优先级" if kind == "operations" else "四、ESG 结论与改进优先级"
    doc.add_heading(tail_title, level=1)
    doc.add_paragraph("总体结论：当前样本覆盖建筑已形成可执行整改清单，建议按“高风险先处置、低风险持续优化”的原则推进。")
    doc.add_paragraph("P0：优先处置异常占比高且 DH 偏低建筑，完成关键设备巡检与故障点闭环。")
    doc.add_paragraph("P1：执行夜间基荷治理、非工作时段自动关停与阈值联动告警。")
    doc.add_paragraph("P2：按月复核预测偏差、建议执行率与节能成效，形成持续改进机制。")
    doc.add_paragraph("验收建议：按周统计“异常占比、工单闭环时效、建议执行率”三项指标，连续 4 周评估整改有效性。")

    doc.save(bio)
    return (
        bio.getvalue(),
        f"{kind}_report.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _build_pdf(ctx: dict[str, Any], kind: str) -> tuple[bytes, str, str]:
    bio = io.BytesIO()
    title = "建筑能源运营优化报告" if kind == "operations" else "建筑能源 ESG 专项报告"
    _register_cn_font()
    font_name = "Chinese" if "Chinese" in pdfmetrics.getRegisteredFontNames() else "Helvetica"

    doc = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
        title=title,
    )
    styles = getSampleStyleSheet()
    st_title = styles["Title"].clone("CNTitle")
    st_title.fontName = font_name
    st_title.fontSize = 18
    st_title.leading = 24
    st_h1 = styles["Heading1"].clone("CNH1")
    st_h1.fontName = font_name
    st_h1.fontSize = 13
    st_h1.leading = 18
    st_h2 = styles["Heading2"].clone("CNH2")
    st_h2.fontName = font_name
    st_h2.fontSize = 11
    st_h2.leading = 16
    st_body = styles["BodyText"].clone("CNBody")
    st_body.fontName = font_name
    st_body.fontSize = 9.5
    st_body.leading = 14
    st_body.spaceAfter = 4

    story: list[Any] = []
    story.append(Paragraph(title, st_title))
    story.append(Spacer(1, 6))
    story.append(Paragraph("项目名称：建筑能源智能管理系统", st_body))
    story.append(Paragraph("文档级别：内部评审版", st_body))
    story.append(Paragraph(f"生成时间：{ctx['generated_at']}", st_body))
    story.append(Paragraph(f"覆盖范围：{ctx['scope_note']}", st_body))
    if ctx.get("building_id"):
        story.append(Paragraph(f"目标建筑：{ctx['building_id']}", st_body))

    story.append(Spacer(1, 8))
    story.append(Paragraph("一、报告说明", st_h1))
    story.append(Paragraph("数据来源：系统小时级能耗数据、异常检测结果与运营规则库。", st_body))
    story.append(Paragraph("分析范围：默认按总电耗选取前 5 栋建筑；若指定建筑 ID 则按单栋输出。", st_body))
    story.append(Paragraph("输出形式：结构化文字结论与整改建议，适用于评审与运维复盘。", st_body))
    story.append(Paragraph("口径说明：异常检测采用 z-score（阈值=3.0）；EWI 反映夜间基荷浪费倾向；DH 反映设备运行健康近似值。", st_body))
    story.append(Paragraph("说明：空间利用相关建议不在本报告输出，室内布局建议请以“孪生与视觉”页面上传建模结果为准。", st_body))

    story.append(Spacer(1, 8))
    story.append(Paragraph("二、建筑清单与核心指标总览", st_h1))
    overview = [["建筑", "风险等级", "样本行数", "累计电耗(kWh)", "异常占比(%)", "EWI", "DH", "未来24h均值(kWh/h)"]]
    for p in ctx["buildings"]:
        period = p.get("period") or {}
        anomaly = p.get("anomaly") or {}
        ind = p.get("indicators") or {}
        overview.append(
            [
                str(p["building_id"]),
                _risk_level(p),
                str(period.get("rows", 0)),
                _fmt((period.get("sums") or {}).get("electricity_kwh")),
                _fmt((anomaly.get("ratio") or 0) * 100),
                _fmt(ind.get("ewi"), 4),
                _fmt(ind.get("dh"), 4),
                _fmt((p.get("forecast") or {}).get("next24_avg_kwh_h")),
            ]
        )
    t = Table(overview, repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#C9CDD4")),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    story.append(t)
    story.append(Paragraph("风险等级判定：综合异常占比、EWI、DH 三项指标，分为高/中/低。", st_body))

    story.append(Spacer(1, 8))
    story.append(Paragraph("三、逐栋异常问题报告与建议", st_h1))
    for idx, p in enumerate(ctx["buildings"], 1):
        period = p.get("period") or {}
        sums = period.get("sums") or {}
        means = period.get("means") or {}
        tr = period.get("time_range") or {}
        anomaly = p.get("anomaly") or {}
        indicators = p.get("indicators") or {}
        fc = p.get("forecast") or {}
        bid = p["building_id"]

        story.append(Paragraph(f"{idx}. 建筑 {bid}", st_h2))
        story.append(
            Paragraph(
                f"监测概况：样本 {period.get('rows', 0)} 行；时间范围 {str(tr.get('min', '-'))[:19]} 至 {str(tr.get('max', '-'))[:19]}；"
                f"累计电耗 {_fmt(sums.get('electricity_kwh'))} kWh；平均小时电耗 {_fmt(means.get('electricity_kwh'))} kWh/h。",
                st_body,
            )
        )
        story.append(
            Paragraph(
                f"指标：异常小时 {anomaly.get('anomaly_hours', 0)} / 总小时 {anomaly.get('total_hours', 0)}，占比 {_fmt((anomaly.get('ratio') or 0) * 100)}%；"
                f"EWI={_fmt(indicators.get('ewi'),4)}，DH={_fmt(indicators.get('dh'),4)}；"
                f"未来24h均值预测={_fmt(fc.get('next24_avg_kwh_h'))} kWh/h（模型：{fc.get('model') or '-'}）。",
                st_body,
            )
        )
        story.append(Paragraph(f"综合风险等级：{_risk_level(p)}。", st_body))
        story.append(Paragraph("异常问题：", st_body))
        for x in p["problems"]:
            story.append(Paragraph(f"• {x}", st_body))

        samples = (anomaly.get("samples") or [])[:5]
        story.append(Paragraph("异常样本（前 5 条）：", st_body))
        if samples:
            s_rows = [["监测时间", "市电(kWh)", "建筑"]]
            for s in samples:
                s_rows.append([str(s.get("monitor_time", "-"))[:19], _fmt(s.get("electricity_kwh")), str(s.get("building_id", bid))])
            st = Table(s_rows, repeatRows=1)
            st.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), font_name),
                        ("FONTSIZE", (0, 0), (-1, -1), 8.2),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F9FAFB")),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D5DD")),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ]
                )
            )
            story.append(st)
        else:
            story.append(Paragraph("当前阈值下无异常样本。", st_body))

        story.append(Paragraph("五段式整改建议（现象-影响-处置-责任人-时限）：", st_body))
        for r in _five_segment_rows(p):
            story.append(Paragraph(f"[{r['priority']}] 现象：{r['phenomenon']}", st_body))
            story.append(Paragraph(f"影响：{r['impact']}", st_body))
            story.append(Paragraph(f"处置：{r['action']}", st_body))
            story.append(Paragraph(f"责任人：{r['owner']}；完成时限：{r['due']}", st_body))

        story.append(Paragraph("整改计划表：", st_body))
        plan = [["优先级", "措施", "责任人", "完成时限"]] + _action_plan_rows(p)
        pt = Table(plan, repeatRows=1)
        pt.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 8.3),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#C9CDD4")),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        story.append(pt)

        if kind == "esg":
            e, s, g = _esg_summary_line(p)
            story.append(Paragraph(f"E（环境）：{e}", st_body))
            story.append(Paragraph(f"S（社会）：{s}", st_body))
            story.append(Paragraph(f"G（治理）：{g}", st_body))
            story.append(Paragraph("ESG专项建议：落实夜间基荷治理与设备巡检闭环，按月形成指标台账。", st_body))
        story.append(Spacer(1, 8))

    story.append(Paragraph("四、结论与实施优先级", st_h1))
    story.append(Paragraph("总体结论：当前样本覆盖建筑已形成可执行整改清单，建议按“高风险先处置、低风险持续优化”的原则推进。", st_body))
    story.append(Paragraph("P0：优先处置异常占比高且 DH 偏低建筑，完成关键设备巡检与故障点闭环。", st_body))
    story.append(Paragraph("P1：执行夜间基荷治理、非工作时段自动关停与阈值联动告警。", st_body))
    story.append(Paragraph("P2：按月复核预测偏差、建议执行率与节能成效，形成持续改进机制。", st_body))
    story.append(Paragraph("验收建议：按周统计“异常占比、工单闭环时效、建议执行率”三项指标，连续 4 周评估整改有效性。", st_body))

    doc.build(story)
    return bio.getvalue(), f"{kind}_report.pdf", "application/pdf"


def build_report_bytes(kind: str, file_format: str, building_id: str | None = None) -> tuple[bytes, str, str]:
    """
    Returns (body, filename, media_type).
    """
    ctx = _report_context(kind, building_id)
    if file_format == "word":
        if not _ensure_docx():
            txt = (
                f"{kind} report fallback: python-docx unavailable. "
                "Please install `python-docx` in the running backend environment."
            ).encode("utf-8")
            return txt, f"{kind}_report_fallback.txt", "text/plain; charset=utf-8"
        return _build_word(ctx, kind)
    if not _ensure_reportlab():
        txt = (
            f"{kind} report fallback: reportlab unavailable. "
            "Please install `reportlab` in the running backend environment."
        ).encode("utf-8")
        return txt, f"{kind}_report_fallback.txt", "text/plain; charset=utf-8"
    return _build_pdf(ctx, kind)


def _wrap_for_canvas(s: str, width: int) -> list[str]:
    if not s:
        return [""]
    out: list[str] = []
    cur = ""
    for ch in s:
        if len(cur) >= width and ch not in (" ", "，", "。"):
            out.append(cur)
            cur = ch
        else:
            cur += ch
    if cur:
        out.append(cur)
    return out if out else [""]


def _register_cn_font() -> None:
    if "Chinese" in pdfmetrics.getRegisteredFontNames():
        return
    candidates: list[tuple[str, int | None]] = [
        (r"C:\Windows\Fonts\msyh.ttc", 0),
        (r"C:\Windows\Fonts\simhei.ttf", None),
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", 0),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", 0),
    ]
    for path, sub in candidates:
        try:
            if sub is None:
                pdfmetrics.registerFont(TTFont("Chinese", path))
            else:
                pdfmetrics.registerFont(TTFont("Chinese", path, subfontIndex=sub))
            return
        except Exception:
            continue
